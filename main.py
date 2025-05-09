import os
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
from datetime import datetime
import re
from collections import Counter
import matplotlib.pyplot as plt
import seaborn as sns
from jinja2 import Template
import sys
import chardet
import webbrowser
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import spacy

# Set UTF-8 encoding for proper handling of Swedish characters
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    # Download the model if it's not installed
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"], check=True)
    nlp = spacy.load("en_core_web_sm")

# Folder containing resumes
RESUME_FOLDER = "resumes/"

# Ensure the resumes folder exists
os.makedirs(RESUME_FOLDER, exist_ok=True)

def verify_document(file_path):
    """
    Verifies a document to ensure it exists, is a PDF or DOCX, and is readable.
    """
    try:
        if not os.path.exists(file_path):
            return f"Error: File does not exist at {file_path}"

        if not (file_path.lower().endswith('.pdf') or file_path.lower().endswith('.docx')):
            return "Error: File is not a supported document type (PDF or DOCX)."

        if file_path.lower().endswith('.pdf'):
            try:
                reader = PdfReader(file_path)
                if len(reader.pages) > 0:
                    return "Document verification successful."
                else:
                    return "Error: PDF appears to be empty."
            except Exception as e:
                return f"Error: Unable to read PDF file. Details: {e}"
        elif file_path.lower().endswith('.docx'):
            try:
                Document(file_path)
                return "Document verification successful."
            except Exception as e:
                return f"Error: Unable to read DOCX file. Details: {e}"

    except Exception as general_error:
        return f"An unexpected error occurred: {general_error}"

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file with improved encoding handling."""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = []
            for page in reader.pages:
                content = page.extract_text()
                if content:
                    # Remove the common header
                    content = re.sub(r'Annette Frohm.*?E-post:', '', content, flags=re.DOTALL)
                    text.append(content)
            
            full_text = '\n'.join(text)
            
            # Get the filename without extension as a fallback name
            filename = os.path.splitext(os.path.basename(file_path))[0]
            name = filename.replace('_', ' ')
            
            # Add the name to the beginning of the text
            full_text = f"Name: {name}\n{full_text}"
            
            return full_text
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from a Word file with improved encoding handling."""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        # Ensure proper encoding of Swedish characters
        if not isinstance(text, str):
            encoding = chardet.detect(text.encode())['encoding']
            text = text.encode().decode(encoding or 'utf-8', errors='ignore')
        return text.strip()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def analyze_resume(text):
    """Analyze resume text using spaCy for better results"""
    doc = nlp(text)
    
    analysis = {
        'skills': extract_skills(doc),
        'education': extract_education(doc),
        'experience': extract_experience(doc),
        'languages': extract_languages(doc),
        'certifications': extract_certifications(doc)
    }
    return analysis

def extract_skills(doc):
    """Extract skills using spaCy's entity recognition and pattern matching"""
    skills = []
    
    # Technical skills patterns
    tech_patterns = [
        'python', 'java', 'javascript', 'html', 'css', 'sql', 'react', 'node.js',
        'aws', 'docker', 'kubernetes', 'git', 'machine learning', 'data analysis'
    ]
    
    # Soft skills patterns
    soft_patterns = [
        'leadership', 'communication', 'teamwork', 'problem solving', 'project management',
        'time management', 'analytical', 'creative', 'detail oriented', 'organization'
    ]
    
    # Extract skills from text
    text_lower = doc.text.lower()
    for skill in tech_patterns + soft_patterns:
        if skill in text_lower:
            skills.append(skill)
    
    # Add skills from entity recognition
    for ent in doc.ents:
        if ent.label_ in ['PRODUCT', 'ORG'] and len(ent.text) > 2:  # Filter out short abbreviations
            skills.append(ent.text)
    
    return list(set(skills))  # Remove duplicates

def extract_education(doc):
    """Extract education using spaCy's entity recognition and pattern matching"""
    education = []
    edu_keywords = ['degree', 'university', 'college', 'school', 'bachelor', 'master', 'phd']
    
    # Look for education-related sentences
    for sent in doc.sents:
        sent_text = sent.text.lower()
        if any(keyword in sent_text for keyword in edu_keywords):
            # Clean up and format the education entry
            clean_text = ' '.join(sent.text.split())
            if clean_text not in education:
                education.append(clean_text)
    
    return education

def extract_experience(doc):
    """Extract work experience using spaCy's entity recognition and date patterns"""
    experience = []
    
    for sent in doc.sents:
        # Look for sentences with dates and organizations
        has_date = any(token.like_num or token.ent_type_ == 'DATE' for token in sent)
        has_org = any(ent.label_ == 'ORG' for ent in sent.ents)
        
        if has_date and has_org:
            clean_text = ' '.join(sent.text.split())
            if clean_text not in experience:
                experience.append(clean_text)
    
    return experience

def extract_languages(doc):
    """Extract languages using spaCy's entity recognition and pattern matching"""
    common_languages = [
        'english', 'spanish', 'french', 'german', 'chinese', 'japanese',
        'arabic', 'russian', 'portuguese', 'italian', 'swedish', 'norwegian'
    ]
    
    languages = []
    text_lower = doc.text.lower()
    
    # Look for common languages
    for lang in common_languages:
        if lang in text_lower:
            languages.append(lang)
    
    # Add languages from entity recognition
    for ent in doc.ents:
        if ent.label_ == 'LANGUAGE':
            languages.append(ent.text)
    
    return list(set(languages))

def extract_certifications(doc):
    """Extract certifications using spaCy's pattern matching"""
    certifications = []
    cert_keywords = ['certified', 'certification', 'certificate', 'license']
    
    for sent in doc.sents:
        sent_text = sent.text.lower()
        if any(keyword in sent_text for keyword in cert_keywords):
            # Clean up and format the certification entry
            clean_text = ' '.join(sent.text.split())
            if clean_text not in certifications:
                certifications.append(clean_text)
    
    return certifications

def calculate_score(analysis):
    """Calculate a comprehensive score based on the analysis"""
    score = 0
    
    # Skills score (up to 30 points)
    skill_score = min(len(analysis['skills']) * 3, 30)
    score += skill_score
    
    # Education score (up to 25 points)
    edu_score = min(len(analysis['education']) * 5, 25)
    score += edu_score
    
    # Experience score (up to 25 points)
    exp_score = min(len(analysis['experience']) * 5, 25)
    score += exp_score
    
    # Languages score (up to 10 points)
    lang_score = min(len(analysis['languages']) * 2, 10)
    score += lang_score
    
    # Certifications score (up to 10 points)
    cert_score = min(len(analysis['certifications']) * 2, 10)
    score += cert_score
    
    return {
        'total': score,
        'breakdown': {
            'skills': skill_score,
            'education': edu_score,
            'experience': exp_score,
            'languages': lang_score,
            'certifications': cert_score
        }
    }

def generate_html_report(results_dict):
    html_content = []
    html_content.append('''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>CV Sammanfattning</title>
            <style>
                :root {
                    --primary-color: #2563eb;
                    --background-color: #f8fafc;
                    --text-color: #1e293b;
                    --border-color: #e2e8f0;
                    --tag-bg: #e0e7ff;
                    --tag-color: #3730a3;
                }
                
                * {
                    margin: 0;
                    padding: 0;
                    box-sizing: border-box;
                }
                
                body {
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    line-height: 1.6;
                    background: var(--background-color);
                    color: var(--text-color);
                }
                
                .container {
                    max-width: 1000px;
                    margin: 2rem auto;
                    padding: 0 1rem;
                }
                
                .header {
                    text-align: center;
                    margin-bottom: 3rem;
                    padding: 2rem 0;
                }
                
                .header h1 {
                    font-size: 2.5rem;
                    color: var(--primary-color);
                    margin-bottom: 1rem;
                }
                
                .cv-list {
                    display: grid;
                    gap: 2rem;
                }
                
                .cv-item {
                    background: white;
                    border-radius: 12px;
                    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
                    padding: 2rem;
                    transition: transform 0.2s ease;
                }
                
                .cv-item:hover {
                    transform: translateY(-2px);
                }
                
                .name {
                    font-size: 1.5rem;
                    font-weight: 600;
                    color: var(--primary-color);
                    margin-bottom: 1.5rem;
                    padding-bottom: 1rem;
                    border-bottom: 2px solid var(--border-color);
                }
                
                .detail {
                    margin: 1.5rem 0;
                }
                
                .label {
                    display: block;
                    font-weight: 600;
                    color: var(--text-color);
                    margin-bottom: 0.5rem;
                    font-size: 0.9rem;
                    text-transform: uppercase;
                    letter-spacing: 0.05em;
                }
                
                .content {
                    background: var(--background-color);
                    padding: 1rem;
                    border-radius: 8px;
                }
                
                ul {
                    list-style: none;
                }
                
                li {
                    margin: 0.5rem 0;
                    padding-left: 1.5rem;
                    position: relative;
                }
                
                li:before {
                    content: "•";
                    color: var(--primary-color);
                    position: absolute;
                    left: 0;
                }
                
                .competence-item {
                    display: inline-block;
                    background: var(--tag-bg);
                    color: var(--tag-color);
                    padding: 0.3rem 0.8rem;
                    margin: 0.25rem;
                    border-radius: 20px;
                    font-size: 0.9rem;
                    font-weight: 500;
                    transition: all 0.2s ease;
                }
                
                .competence-item:hover {
                    transform: scale(1.05);
                }
                
                @media (max-width: 768px) {
                    .container {
                        margin: 1rem auto;
                    }
                    
                    .cv-item {
                        padding: 1.5rem;
                    }
                    
                    .header h1 {
                        font-size: 2rem;
                    }
                }
            </style>
        </head>
        <body>
            <div class="container">
                <header class="header">
                    <h1>CV Sammanfattning</h1>
                </header>
                <div class="cv-list">
    ''')

    for name, result in results_dict.items():
        html_content.append(f'''
            <div class="cv-item">
                <div class="name">{name}</div>
                
                <div class="detail">
                    <span class="label">Färdigheter</span>
                    <div class="content">
                        {''.join(f"<span class='competence-item'>{skill}</span>" for skill in result['skills']) if result['skills'] else 'Inga färdigheter angivna'}
                    </div>
                </div>

                <div class="detail">
                    <span class="label">Utbildning</span>
                    <div class="content">
                        <ul>
                            {''.join(f"<li>{edu}</li>" for edu in result['education']) if result['education'] else '<li>Ingen utbildning angiven</li>'}
                        </ul>
                    </div>
                </div>

                <div class="detail">
                    <span class="label">Arbetslivserfarenhet</span>
                    <div class="content">
                        <ul>
                            {''.join(f"<li>{exp}</li>" for exp in result['experience']) if result['experience'] else '<li>Ingen arbetslivserfarenhet angiven</li>'}
                        </ul>
                    </div>
                </div>

                <div class="detail">
                    <span class="label">Språkkunskaper</span>
                    <div class="content">
                        <ul>
                            {''.join(f"<li>{lang}</li>" for lang in result['languages']) if result['languages'] else '<li>Inga språk angivna</li>'}
                        </ul>
                    </div>
                </div>

                <div class="detail">
                    <span class="label">Certifikat</span>
                    <div class="content">
                        <ul>
                            {''.join(f"<li>{cert}</li>" for cert in result['certifications']) if result['certifications'] else '<li>Inga certifikat angivna</li>'}
                        </ul>
                    </div>
                </div>
            </div>
        ''')

    html_content.append('''
                </div>
            </div>
        </body>
        </html>
    ''')

    return ''.join(html_content)

def send_email_report(recipient_email, sender_email, sender_password):
    """Send the resume report via email"""
    try:
        # Create the email message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = 'Resume Analysis Report'

        # Add body text
        body = 'Please find attached the resume analysis report.'
        msg.attach(MIMEText(body, 'plain'))

        # Attach the HTML report
        with open('resume_report.html', 'r', encoding='utf-8') as f:
            attachment = MIMEText(f.read(), 'html')
            attachment.add_header('Content-Disposition', 'attachment', filename='resume_report.html')
            msg.attach(attachment)

        # Connect to Gmail's SMTP server
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)

        # Send the email
        server.send_message(msg)
        server.quit()
        print(f'Report sent successfully to {recipient_email}')
        return True
    except Exception as e:
        print(f'Error sending email: {str(e)}')
        return False

def main():
    # Process all resumes in the folder
    results = []
    for filename in os.listdir(RESUME_FOLDER):
        if filename.startswith('.'):  # Skip hidden files like .DS_Store
            continue
            
        file_path = os.path.join(RESUME_FOLDER, filename)
        try:
            if verify_document(file_path):
                if file_path.lower().endswith('.pdf'):
                    content = extract_text_from_pdf(file_path)
                else:
                    content = extract_text_from_docx(file_path)
                
                result = analyze_resume(content)
                if result:
                    # Add filename as name if not found in content
                    if 'Name' not in result or not result['Name']:
                        result['Name'] = os.path.splitext(filename)[0].replace('_', ' ')
                    results.append(result)
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")

    # Convert results to dictionary
    results_dict = {result['Name']: result for result in results}
    
    # Calculate scores
    scores = [calculate_score(result) for result in results_dict.values()]
    
    # Generate and save the HTML report
    html_content = generate_html_report(results_dict)
    
    # Save the report
    report_file = 'resume_report.html'
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # Save results to CSV
    if results:
        df = pd.DataFrame(results)
        df.to_csv('cv_report.csv', index=False, encoding='utf-8')
        print(f"Successfully processed {len(results)} resumes")

if __name__ == "__main__":
    main()
    # Open the generated report in the default web browser
    report_path = os.path.abspath('resume_report.html')
    print(f"Opening report: {report_path}")
    webbrowser.open('file://' + report_path)
